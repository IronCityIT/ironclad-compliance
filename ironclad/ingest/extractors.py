"""Text extraction from evidence files.

Re-housed from the original scripts/assess_controls.py, with the behaviour that
mattered kept and the parts that hid problems fixed:

  * the original swallowed every extraction failure into a bare except and
    returned a placeholder string, so a corrupt PDF looked exactly like a PDF
    with no relevant content. Failures are now reported to the caller.
  * the original truncated at 5000 characters for every format. That cap exists
    to bound memory across a large evidence set, so it stays, but it is now a
    named constant and the truncation is recorded rather than silent.

Binary format support is optional. An engine that cannot read a PDF should say
so on that one artifact and keep assessing the rest, not fail the run.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

# Bounds memory across a large evidence set. Control matching keys on evidence
# type and section headings, which sit near the top of a policy document, so the
# cap costs little recall.
MAX_CHARS = 20_000
MAX_PDF_PAGES = 20
MAX_SHEETS = 5
MAX_SHEET_ROWS = 200

TEXT_SUFFIXES = frozenset({".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml", ".html"})


@dataclass
class Extraction:
    """What came out of one file, and whether anything went wrong getting it."""

    text: str = ""
    truncated: bool = False
    error: str = ""

    @property
    def ok(self) -> bool:
        return not self.error


def supported_extensions() -> list[str]:
    """Every suffix the extractor will attempt, for the ingestion docs and UI."""
    return sorted(TEXT_SUFFIXES | {".pdf", ".docx", ".xlsx", ".xls"})


def _clip(text: str) -> Extraction:
    if len(text) > MAX_CHARS:
        return Extraction(text=text[:MAX_CHARS], truncated=True)
    return Extraction(text=text)


def _extract_pdf(path: Path) -> Extraction:
    try:
        import PyPDF2  # noqa: PLC0415 — optional dependency, imported on demand
    except ImportError:
        return Extraction(error="PDF support is not installed (PyPDF2)")

    try:
        with path.open("rb") as handle:
            reader = PyPDF2.PdfReader(handle)
            pages = [(page.extract_text() or "") for page in reader.pages[:MAX_PDF_PAGES]]
        return _clip("\n".join(pages))
    except Exception as exc:  # noqa: BLE001 — any parser fault is reported, not raised
        return Extraction(error=f"could not read PDF: {exc}")


def _extract_docx(path: Path) -> Extraction:
    try:
        from docx import Document  # noqa: PLC0415 — optional dependency
    except ImportError:
        return Extraction(error="DOCX support is not installed (python-docx)")

    try:
        document = Document(str(path))
        parts = [p.text for p in document.paragraphs]
        # Control evidence often lives in tables (access matrices, review logs),
        # which the original extractor skipped entirely.
        for table in document.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        return _clip("\n".join(parts))
    except Exception as exc:  # noqa: BLE001
        return Extraction(error=f"could not read DOCX: {exc}")


def _extract_xlsx(path: Path) -> Extraction:
    try:
        import openpyxl  # noqa: PLC0415 — optional dependency
    except ImportError:
        return Extraction(error="spreadsheet support is not installed (openpyxl)")

    try:
        workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in workbook.worksheets[:MAX_SHEETS]:
            lines.append(str(sheet.title))
            for row in sheet.iter_rows(max_row=MAX_SHEET_ROWS, values_only=True):
                lines.append(" ".join(str(cell) for cell in row if cell is not None))
        workbook.close()
        return _clip("\n".join(lines))
    except Exception as exc:  # noqa: BLE001
        return Extraction(error=f"could not read spreadsheet: {exc}")


def _extract_text_file(path: Path) -> Extraction:
    try:
        return _clip(path.read_text(encoding="utf-8", errors="replace"))
    except OSError as exc:
        return Extraction(error=f"could not read file: {exc}")


def extract_text(path: Path) -> Extraction:
    """Pull searchable text out of one evidence file.

    Never raises for an unreadable file — an assessment must not die because one
    artifact in a hundred is corrupt. The fault travels on the Extraction and
    ends up as a note on the affected control.
    """
    if not path.exists():
        return Extraction(error="file not found")
    if not path.is_file():
        return Extraction(error="not a file")

    suffix = path.suffix.lower()
    if suffix in TEXT_SUFFIXES:
        return _extract_text_file(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in (".xlsx", ".xls"):
        return _extract_xlsx(path)
    return Extraction(error=f"unsupported evidence format {suffix or '(none)'}")
